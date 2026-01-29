from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from pathlib import Path


def add_code_block(doc, code_text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def build_document(programs, output_path="Assignment.docx"):
    doc = Document()

    for idx, program in enumerate(programs):
        # ---- Title ----
        doc.add_heading(program["name"], level=1)

        # ---- Language ----
        lang_para = doc.add_paragraph()
        lang_para.add_run("Language: ").bold = True
        lang_para.add_run(program["language"].upper())

        # ---- Source Code ----
        doc.add_paragraph("\nSource Code:", style="List Bullet")

        for src in program["source_files"]:
            doc.add_paragraph(f"\nFile: {src.name}", style="Intense Quote")

            code_text = src.read_text(encoding="utf-8", errors="ignore")
            add_code_block(doc, code_text)

        # ---- Execution Output ----
        doc.add_paragraph("\nExecution Output:", style="List Bullet")

        for img_path in program.get("output_images", []):
            doc.add_picture(str(img_path))

        # ---- Page break (except last) ----
        if idx < len(programs) - 1:
            doc.add_page_break()

    doc.save(output_path)
